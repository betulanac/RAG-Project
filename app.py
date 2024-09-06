from langchain.retrievers.contextual_compression import ContextualCompressionRetriever
from langchain_cohere import CohereRerank
import streamlit as st
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from langchain.vectorstores import FAISS
from langchain_community.embeddings import CohereEmbeddings
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain
from langchain_openai import AzureChatOpenAI
import logging
from langchain.docstore.document import Document
import os
from docx import Document as DocxDocument 
import requests
from bs4 import BeautifulSoup
from langchain.text_splitter import RecursiveCharacterTextSplitter


logging.basicConfig(level=logging.DEBUG)

load_dotenv()
api_key = os.getenv("AZURE_OPENAI_API_KEY")
api_endpoint = os.getenv('AZURE_OPENAI_ENDPOINT')

client = AzureChatOpenAI(
    azure_endpoint=api_endpoint,
    api_key=api_key,
    api_version="2024-02-15-preview"
)

cohere_embeddings = CohereEmbeddings(cohere_api_key=os.getenv("COHERE_API_KEY"),
                                     model="embed-english-light-v3.0",user_agent="agent")

compressor = CohereRerank(cohere_api_key=os.getenv("COHERE_API_KEY"), model="rerank-english-v3.0")

# Function to retrieve content from the website
def get_text_from_web(url):
    try:
        response = requests.get(url)
        if response.status_code == 200:
            soup = BeautifulSoup(response.content, 'html.parser')
            text = soup.get_text(separator="\n")
            return text
        else:
            st.warning(f"Failed to retrieve the website. Status code: {response.status_code}")
            return None
    except Exception as e:
        st.error(f"An error occurred while fetching the website: {e}")
        return None
    
#ContextualCompressionRetriever is a compression retriever. 
# Its purpose is to compress returned documents and extract more relevant information.
def get_compression_retriever(vector_store):
    base_retriever = vector_store.as_retriever()
    compression_retriever = ContextualCompressionRetriever(
        base_compressor=compressor, #Compresses or reorders documents to create a more focused list of results.
        base_retriever=base_retriever
    )
    return compression_retriever

# Text import function from files
def get_text_from_files(files):
    text = ""
    for file in files:
        if file.name.endswith(".pdf"):
            reader = PdfReader(file)
            for page in reader.pages:
                text += page.extract_text()
        elif file.name.endswith(".docx"):
            doc = DocxDocument(file)
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif file.name.endswith(".txt"):
            text += file.read().decode("utf-8") + "\n"
        else:
            st.warning(f"Unsupported file type: {file.name}")
    return text

def get_text_chunks(text):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )
    chunks = text_splitter.split_text(text)
    return chunks

def get_vectorstore(text_chunks):
    documents = [Document(page_content=chunk) for chunk in text_chunks]
    
    vector_store = FAISS.from_documents(documents=documents, embedding=cohere_embeddings)
    return vector_store

def get_conversation_chain(vector_store):
    llm_model = "gpt-4o"
    
    llm = AzureChatOpenAI(temperature=0.0, model=llm_model, api_version="2024-02-15-preview")
    
    memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)
    conversation_chain = ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=vector_store.as_retriever(),
        memory=memory
    )
    
    st.session_state.vector_store = vector_store
    st.session_state.conversation = conversation_chain
    st.write("Conversation chain created and stored in session state.") 
    return conversation_chain

def handle_userinput(user_question):
    if st.session_state.conversation:
        try:
            vector_store = st.session_state.vector_store
            compression_retriever = get_compression_retriever(vector_store)

            # Similarity search with compressed documents
            similar_docs_with_scores = compression_retriever.get_relevant_documents(user_question)
            
            if similar_docs_with_scores:
                doc = similar_docs_with_scores[0]
                response_text = doc.page_content
                #st.write(f"Found compressed document: {response_text}")
            else:
                st.write("No similar documents found.")
                
            response = st.session_state.conversation({'question': user_question})
            st.session_state.chat_history = response['chat_history']

            for i, message in enumerate(st.session_state.chat_history):
                if i % 2 == 0:
                    st.chat_message("user").write(message.content)
                else:
                    st.chat_message("assistant").write(message.content)
        
        except Exception as e:
            st.error(f"An error occurred: {e}")
            logging.error(f"An error occurred: {e}", exc_info=True)
    else:
        st.warning("No conversation chain found.")


def main():
    if "conversation" not in st.session_state:
        st.session_state.conversation = None

    if "chat_history" not in st.session_state:
        st.session_state.chat_history = None

    st.set_page_config(page_title="RAG")

    st.header("RAG: Chat With Your Document")

    # Getting input from user
    user_question = st.chat_input("Ask a question") 

    if user_question:
        if st.session_state.conversation:
            try:
               handle_userinput(user_question)
            except Exception as e:
                st.error(f"An error occurred: {e}")
                logging.error(f"An error occurred: {e}", exc_info=True)
        else:
            st.warning("Please upload and process a document first.")

    with st.sidebar:
        st.subheader("Documents")
        files = st.file_uploader("Upload your PDF, DOCX, or TXT files", accept_multiple_files=True, type=["pdf", "docx", "txt"])
        url = st.text_input("Or enter a URL to fetch content from the web")

        if st.button("Process"):
            if files or url:
                try:
                    with st.spinner("Processing"):
                        raw_text = ""

                        # Read files
                        if files:
                            raw_text += get_text_from_files(files)

                        # Get content from URL
                        if url:
                            web_text = get_text_from_web(url)
                            if web_text:
                                raw_text += web_text

                        # Break up text and process it
                        text_chunks = get_text_chunks(raw_text)

                        vector_store = get_vectorstore(text_chunks)
                        st.session_state.vector_store = vector_store
                        st.session_state.conversation = get_conversation_chain(vector_store) 
                except Exception as e:
                    st.error(f"An error occurred during processing: {e}")
                    logging.error(f"An error occurred during processing: {e}", exc_info=True)
            else:
                st.warning("No documents or URLs provided.")

if __name__ == "__main__":
    main()